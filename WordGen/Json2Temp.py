from docx import Document
import logging
import json
import os

__all__ = ["wordgen_json2temp"]

WORDGEN_TABLE_STYLE = "ListenAI_Table"


def wordgen_json2temp(doc):
    with open(os.path.join(os.getcwd(), "__info", "ex2js.json"), "r") as fr:
        handle = json.load(fr)

    logging.info("Generate template...")
    for item in handle:
        jinjia_key = "{{ %s_Name }}" % item["module"]
        doc.add_heading(jinjia_key, level=2)

        # jinjia_key = "{{ %s_Address}}" %  item["module"]
        # doc.add_paragraph(jinjia_key)

        for in_item in item["registers"]:
            jinjia_key = "{{ %s_%s_Name }}" % (item["module"], in_item["register"])
            doc.add_heading(jinjia_key, level=3)

            jinjia_key = "{{ %s_%s_Offset }}" % (item["module"], in_item["register"])
            doc.add_paragraph("Offset: %s" % jinjia_key)

            # generate table
            # merge
            table = doc.add_table(4, 3, WORDGEN_TABLE_STYLE)
            a = table.rows[1].cells[0]
            b = table.rows[1].cells[2]
            a.merge(b)

            a = table.rows[3].cells[0]
            b = table.rows[3].cells[2]
            a.merge(b)

            jinjia_key = "{%tc for col in Module_Register_Col_Labels %}"
            table.cell(0, 0).text = jinjia_key
            table.cell(0, 1).text = "{{ col }}"
            table.cell(0, 2).text = "{%tc endfor %}"

            jinjia_key = "{%%tr for item in %s_%s_contents %%}" % (item["module"], in_item["register"])
            table.cell(1, 0).text = jinjia_key
            table.cell(2, 0).text = "{%tc for col in item %}"
            table.cell(2, 1).text = "{{ col }}"
            table.cell(2, 2).text = "{%tc endfor %}"
            table.cell(3, 0).text = "{%tr endfor %}"

    logging.info("Generate complete")


if __name__ == "__main__":
    doc = Document("../__info/Empty_Ref.docx")
    wordgen_json2temp(doc)
    doc.save("../.Template/WordGen.docx")
